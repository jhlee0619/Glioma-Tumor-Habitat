"""Train classifiers on quantization ratios to predict molecular/prognostic labels.

Targets: who_grade, idh, _1p19q, ki_67.
Methods: AE, dVAE, metric.

- ki_67은 >10:1, <=10:0 이진화.
- Classifier: RandomForestClassifier.
- Metrics: Accuracy, Macro-F1, AUROC(멀티클래스는 macro-OVR), OvR-macro Sensitivity/Specificity.
- 동일 테스트셋에서 두 모델 AUC 비교는 DeLong test(쌍체 DeLong)로 수행.
  * Binary: z = (AUC1-AUC2)/sqrt(Var1+Var2-2*Cov12), p=2[1-Φ(|z|)].
  * Multiclass: OvR 점수를 (공통 유효 클래스에 한해) 스택해 단일 이진 문제로 변환 후 쌍체 DeLong 1회 적용(관행적 근사).
- test_upenn은 평가하지 않음.
"""

import argparse
from pathlib import Path
from typing import Dict, List, Tuple

import re
import joblib
import numpy as np
import pandas as pd
from docx import Document
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (
    accuracy_score, auc, classification_report, f1_score,
    roc_curve, confusion_matrix, roc_auc_score
)

# ---------------------------
# Utilities
# ---------------------------

def sanitize_text(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", value)

def save_metrics_word(metrics_df: pd.DataFrame, output_path: Path) -> None:
    """테스트 결과만 포함. target 1차, method(AE,dVAE,metric) 2차 정렬. p-value 규칙 적용. AUC, Accuracy, Sensitivity, Specificity 모두 CI 표기."""
    document = Document()
    document.add_heading("Classification Metrics", level=1)

    table = document.add_table(rows=1, cols=7)
    header = table.rows[0].cells
    header[0].text = "Target"
    header[1].text = "Method"
    header[2].text = "Accuracy"
    header[3].text = "Sensitivity"
    header[4].text = "Specificity"
    header[5].text = "AUC"
    header[6].text = "p value"

    df = metrics_df[metrics_df["split"] == "test"].copy()
    order = pd.CategoricalDtype(categories=["AE", "dVAE", "metric"], ordered=True)
    df["feature_set"] = df["feature_set"].astype(order)
    df = df.sort_values(["target", "feature_set"]).reset_index(drop=True)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["target"])
        cells[1].text = str(row["feature_set"])
        
        # Accuracy with CI
        if not pd.isna(row["accuracy"]):
            lo = row.get("accuracy_ci_lower", np.nan)
            hi = row.get("accuracy_ci_upper", np.nan)
            cells[2].text = (
                f"{row['accuracy']:.4f}" if pd.isna(lo) or pd.isna(hi)
                else f"{row['accuracy']:.4f} ({lo:.4f}, {hi:.4f})"
            )
        else:
            cells[2].text = "N/A"
        
        # Sensitivity with CI
        if not pd.isna(row.get("sensitivity", np.nan)):
            lo = row.get("sensitivity_ci_lower", np.nan)
            hi = row.get("sensitivity_ci_upper", np.nan)
            cells[3].text = (
                f"{row['sensitivity']:.4f}" if pd.isna(lo) or pd.isna(hi)
                else f"{row['sensitivity']:.4f} ({lo:.4f}, {hi:.4f})"
            )
        else:
            cells[3].text = "N/A"
        
        # Specificity with CI
        if not pd.isna(row.get("specificity", np.nan)):
            lo = row.get("specificity_ci_lower", np.nan)
            hi = row.get("specificity_ci_upper", np.nan)
            cells[4].text = (
                f"{row['specificity']:.4f}" if pd.isna(lo) or pd.isna(hi)
                else f"{row['specificity']:.4f} ({lo:.4f}, {hi:.4f})"
            )
        else:
            cells[4].text = "N/A"
        
        # AUC with CI
        if not pd.isna(row["auroc"]):
            lo = row.get("auroc_ci_lower", np.nan)
            hi = row.get("auroc_ci_upper", np.nan)
            cells[5].text = (
                f"{row['auroc']:.4f}" if pd.isna(lo) or pd.isna(hi)
                else f"{row['auroc']:.4f} ({lo:.4f}, {hi:.4f})"
            )
        else:
            cells[5].text = "N/A"

        meth = str(row["feature_set"])
        if meth == "AE":
            val = row.get("p_value_dvae_vs_ae", np.nan)
            cells[6].text = f"{val:.4f}" if not pd.isna(val) else "N/A"
        elif meth == "metric":
            val = row.get("p_value_dvae_vs_metric", np.nan)
            cells[6].text = f"{val:.4f}" if not pd.isna(val) else "N/A"
        else:
            cells[6].text = "N/A"

    document.save(output_path)
    print(f"Saved metrics DOCX to {output_path}")



def save_model_specific_results(metrics_df: pd.DataFrame, output_dir: Path) -> None:
    """모델별로 분류 지표를 저장합니다."""
    # 모델별로 그룹화
    for feature_set in metrics_df['feature_set'].unique():
        model_dir = output_dir / f"model_{feature_set}"
        model_dir.mkdir(parents=True, exist_ok=True)
        
        # 해당 모델의 데이터만 필터링
        model_metrics = metrics_df[metrics_df['feature_set'] == feature_set].copy()
        
        # CSV 저장
        csv_path = model_dir / "classification_metrics.csv"
        model_metrics.to_csv(csv_path, index=False)
        print(f"Saved {feature_set} metrics CSV to {csv_path}")
        
        # Word 문서 저장
        docx_path = model_dir / "classification_metrics.docx"
        save_metrics_word(model_metrics, docx_path)
        print(f"Saved {feature_set} metrics DOCX to {docx_path}")

# ---------------------------
# Config
# ---------------------------

METHODS = ["AE", "dVAE", "metric"]
TARGETS = {
    "who_grade": [2, 3, 4],
    "idh": [0, 1],
    "_1p19q": [0, 1],
    "ki_67": [0, 1],
}

# ---------------------------
# Data handling
# ---------------------------

def load_dataset(csv_path: str) -> pd.DataFrame:
    df = pd.read_csv(csv_path)
    if "phase" not in df.columns:
        raise ValueError("Dataset must contain a 'phase' column to separate train/test subsets.")
    return df

def get_method_feature_columns(df: pd.DataFrame, method: str) -> List[str]:
    prefix = f"{method}_ratio_"
    return sorted([c for c in df.columns if c.startswith(prefix)])

def build_feature_sets(df: pd.DataFrame, methods: List[str]) -> Dict[str, List[str]]:
    feature_sets: Dict[str, List[str]] = {}
    for method in methods:
        cols = get_method_feature_columns(df, method)
        if cols:
            feature_sets[method] = cols
    return feature_sets

def split_by_phase(df: pd.DataFrame, phase: str) -> pd.DataFrame:
    mask = df["phase"].astype(str).str.lower() == phase
    return df.loc[mask].copy()

def prepare_ki67_data(df: pd.DataFrame) -> pd.DataFrame:
    tdf = df.dropna(subset=['ki_67']).copy()
    tdf['ki_67'] = (tdf['ki_67'] > 10).astype(int)
    return tdf

def prepare_target_data(df: pd.DataFrame, target: str) -> pd.DataFrame:
    if target == "ki_67":
        return prepare_ki67_data(df)
    tdf = df.copy()
    return tdf.dropna(subset=[target])

# ---------------------------
# Model
# ---------------------------

def train_model(X: np.ndarray, y: np.ndarray) -> RandomForestClassifier:
    clf = RandomForestClassifier(
        n_estimators=50, max_depth=2, min_samples_leaf=40, min_samples_split=10,
        max_features=4, bootstrap=True, class_weight="balanced_subsample",
        random_state=42, n_jobs=16,
    )
    clf.fit(X, y)
    return clf

# ---------------------------
# Metrics helpers
# ---------------------------

def compute_ovr_sensitivity_specificity(
    y_true: np.ndarray, y_pred: np.ndarray, classes: np.ndarray
) -> Dict[str, float]:
    if y_true.size == 0:
        return {"sensitivity": np.nan, "specificity": np.nan}
    present = [c for c in classes if (y_true == c).any()]
    if not present:
        return {"sensitivity": np.nan, "specificity": np.nan}
    
    sens, spec = [], []
    for c in present:
        yt = (y_true == c).astype(int)
        yp = (y_pred == c).astype(int)
        cm = confusion_matrix(yt, yp, labels=[0, 1])
        if cm.shape != (2, 2):
            continue
        tn, fp, fn, tp = cm.ravel()
        sens.append(tp / (tp + fn) if (tp + fn) > 0 else np.nan)
        spec.append(tn / (tn + fp) if (tn + fp) > 0 else np.nan)

    sens = [x for x in sens if not np.isnan(x)]
    spec = [x for x in spec if not np.isnan(x)]
    return {
        "sensitivity": float(np.mean(sens)) if sens else np.nan,
        "specificity": float(np.mean(spec)) if spec else np.nan,
    }

# ---------------------------
# DeLong (binary, correlated)
# ---------------------------

def _midrank(x: np.ndarray) -> np.ndarray:
    J = x.size
    order = np.argsort(x)
    xs = x[order]
    r = np.zeros(J, float)
    i = 0
    while i < J:
        j = i
        while j < J and xs[j] == xs[i]:
            j += 1
        r[i:j] = 0.5 * (i + j - 1) + 1.0
        i = j
    out = np.empty(J, float)
    out[order] = r
    return out

def _fast_delong(preds_sorted_T: np.ndarray, n_pos: int) -> Tuple[np.ndarray, np.ndarray]:
    m = n_pos
    n = preds_sorted_T.shape[1] - m
    if m == 0 or n == 0:
        k = preds_sorted_T.shape[0]
        return np.full(k, np.nan), np.full((k, k), np.nan)
    k = preds_sorted_T.shape[0]
    pos = preds_sorted_T[:, :m]
    neg = preds_sorted_T[:, m:]
    tx = np.empty((k, m)); ty = np.empty((k, n)); tz = np.empty((k, m+n))
    for r in range(k):
        tx[r] = _midrank(pos[r]); ty[r] = _midrank(neg[r]); tz[r] = _midrank(preds_sorted_T[r])
    aucs = (tz[:, :m].sum(axis=1) / m - (m + 1) / 2.0) / n
    v01 = (tz[:, :m] - tx) / n
    v10 = 1.0 - (tz[:, m:] - ty) / m
    sx = np.atleast_2d(np.cov(v10, bias=True))
    sy = np.atleast_2d(np.cov(v01, bias=True))
    cov = sx / m + sy / n
    return aucs, cov

def _delong_auc_var(y_true: np.ndarray, y_score: np.ndarray) -> Tuple[float, float]:
    y_true = np.asarray(y_true, int)
    y_score = np.asarray(y_score, float)
    if not set(np.unique(y_true)).issubset({0, 1}):
        return np.nan, np.nan
    order = np.argsort(-y_true)
    ys = y_true[order]; n_pos = int(ys.sum())
    if n_pos == 0 or n_pos == len(ys):
        return np.nan, np.nan
    preds_sorted_T = y_score[order][None, :]
    aucs, cov = _fast_delong(preds_sorted_T, n_pos)
    return float(aucs[0]), float(np.atleast_2d(cov)[0, 0])

def delong_pvalue_correlated_binary(y_true: np.ndarray, s1: np.ndarray, s2: np.ndarray) -> float:
    """쌍체 DeLong p-value for binary ROC AUCs."""
    from scipy.stats import norm
    auc1, v1 = _delong_auc_var(y_true, s1)
    auc2, v2 = _delong_auc_var(y_true, s2)
    if any(np.isnan([auc1, v1, auc2, v2])):
        return np.nan
    order = np.argsort(-y_true)
    ys = y_true[order]; n_pos = int(ys.sum())
    preds_sorted_T = np.vstack([s1[order], s2[order]])
    _, cov = _fast_delong(preds_sorted_T, n_pos)
    if np.any(~np.isfinite(cov)):
        return np.nan
    cov12 = float(np.atleast_2d(cov)[0, 1])
    var_diff = v1 + v2 - 2.0 * cov12
    if not np.isfinite(var_diff) or var_diff <= 0:
        return np.nan
    z = (auc1 - auc2) / np.sqrt(var_diff)
    p = 2.0 * (1.0 - norm.cdf(abs(z)))
    return float(p)

# ---------------------------
# Multiclass: OvR 스택 후 DeLong(관행적 근사, 정합성 보장)
# ---------------------------

def build_ovr_stack_pairwise(y: np.ndarray, proba1: np.ndarray, proba2: np.ndarray, classes: np.ndarray) -> Tuple[np.ndarray, np.ndarray, np.ndarray]:
    """
    두 모델에서 모두 유효한 OvR 이진문제(양/음성 존재 ∧ 점수가 상수 아님)만 선택.
    같은 순서로 y, s1, s2를 스택해 반환.
    """
    y = np.asarray(y)
    P1 = np.asarray(proba1)
    P2 = np.asarray(proba2)
    classes = np.asarray(classes)
    y_list, s1_list, s2_list = [], [], []
    for j, c in enumerate(classes):
        yb = (y == c).astype(int)
        if yb.min() == yb.max():
            continue
        s1 = P1[:, j]; s2 = P2[:, j]
        if np.allclose(s1.max(), s1.min()) or np.allclose(s2.max(), s2.min()):
            continue
        y_list.append(yb); s1_list.append(s1); s2_list.append(s2)
    if not y_list:
        return np.array([], dtype=int), np.array([], dtype=float), np.array([], dtype=float)
    y_stack = np.concatenate(y_list, axis=0)
    s1_stack = np.concatenate(s1_list, axis=0)
    s2_stack = np.concatenate(s2_list, axis=0)
    return y_stack, s1_stack, s2_stack

def delong_pvalue_correlated_multiclass(y: np.ndarray, proba1: np.ndarray, proba2: np.ndarray, classes: np.ndarray) -> float:
    """OvR 스택 기반 멀티클래스 AUC 비교의 쌍체 DeLong 근사 p-value."""
    y_stack, s1_stack, s2_stack = build_ovr_stack_pairwise(y, proba1, proba2, classes)
    if y_stack.size == 0:
        return np.nan
    return delong_pvalue_correlated_binary(y_stack, s1_stack, s2_stack)

# ---------------------------
# AUC bootstrap (binary CI)
# ---------------------------

def bootstrap_binary_auc(y_true: np.ndarray, pos_scores: np.ndarray, n_bootstraps: int = 5000, random_state: int = 42):
    if len(np.unique(y_true)) != 2:
        return None
    rng = np.random.default_rng(random_state)
    vals = []
    y_true = np.asarray(y_true); s = np.asarray(pos_scores)
    for _ in range(n_bootstraps):
        idx = rng.integers(0, len(y_true), len(y_true))
        yb = y_true[idx]; sb = s[idx]
        if yb.min() == yb.max():
            continue
        fpr, tpr, _ = roc_curve(yb, sb, pos_label=1)
        vals.append(auc(fpr, tpr))
    if not vals:
        return None
    arr = np.array(vals)
    return {
        "mean": float(arr.mean()),
        "std": float(arr.std(ddof=1) if len(arr) > 1 else 0.0),
        "ci_lower": float(np.percentile(arr, 2.5)),
        "ci_upper": float(np.percentile(arr, 97.5)),
        "samples": int(len(arr)),
    }

# ---------------------------
# Macro-AUC bootstrap (multiclass CI)
# ---------------------------

def bootstrap_macro_auc(
    y_true: np.ndarray, proba: np.ndarray, class_labels: np.ndarray,
    n_bootstraps: int = 5000, random_state: int = 42,
) -> Dict[str, float]:
    if y_true.size == 0:
        return {"mean": np.nan, "std": np.nan, "ci_lower": np.nan, "ci_upper": np.nan, "samples": 0}
    rng = np.random.default_rng(random_state)
    scores = []
    for _ in range(n_bootstraps):
        idx = rng.integers(0, len(y_true), len(y_true))
        yb = y_true[idx]; pb = proba[idx]
        aucs = []
        for j, cls in enumerate(class_labels):
            ybin = (yb == cls).astype(int)
            if ybin.min() == ybin.max():
                aucs = []; break
            fpr, tpr, _ = roc_curve(ybin, pb[:, j])
            aucs.append(auc(fpr, tpr))
        if aucs:
            scores.append(float(np.mean(aucs)))
    if not scores:
        return {"mean": np.nan, "std": np.nan, "ci_lower": np.nan, "ci_upper": np.nan, "samples": 0}
    s = np.array(scores)
    return {
        "mean": float(s.mean()),
        "std": float(s.std(ddof=1) if len(s) > 1 else 0.0),
        "ci_lower": float(np.percentile(s, 2.5)),
        "ci_upper": float(np.percentile(s, 97.5)),
        "samples": int(len(s)),
    }

def bootstrap_metric_cis(y_true: np.ndarray, y_pred: np.ndarray, classes: np.ndarray,
                         n_bootstraps: int = 5000, random_state: int = 42):
    """
    테스트셋 (y_true, y_pred)로부터 Accuracy, OvR-macro Sensitivity/Specificity의 95% CI 산출.
    모델 재학습/재예측 없이 인덱스 부트스트랩. 비정상 표본(양/음성 단일 클래스 등)은 스킵.
    """
    rng = np.random.default_rng(random_state)
    acc_vals, sens_vals, spec_vals = [], [], []

    y_true = np.asarray(y_true)
    y_pred = np.asarray(y_pred)

    n = len(y_true)
    for _ in range(n_bootstraps):
        idx = rng.integers(0, n, n)
        yt = y_true[idx]; yp = y_pred[idx]
        # 클래스가 하나로 붕괴되면 OvR metric이 정의되지 않으므로 제외
        if len(np.unique(yt)) < 2:
            continue

        acc_vals.append(accuracy_score(yt, yp))
        ovr = compute_ovr_sensitivity_specificity(yt, yp, classes)
        if not np.isnan(ovr["sensitivity"]):
            sens_vals.append(ovr["sensitivity"])
        if not np.isnan(ovr["specificity"]):
            spec_vals.append(ovr["specificity"])

    def _ci(a):
        if not a:
            return (np.nan, np.nan, 0)
        a = np.array(a, float)
        return (float(np.percentile(a, 2.5)),
                float(np.percentile(a, 97.5)),
                int(len(a)))

    acc_lo, acc_hi, acc_n = _ci(acc_vals)
    sen_lo, sen_hi, sen_n = _ci(sens_vals)
    spe_lo, spe_hi, spe_n = _ci(spec_vals)

    return {
        "accuracy_ci_lower": acc_lo, "accuracy_ci_upper": acc_hi, "accuracy_boot_n": acc_n,
        "sensitivity_ci_lower": sen_lo, "sensitivity_ci_upper": sen_hi, "sensitivity_boot_n": sen_n,
        "specificity_ci_lower": spe_lo, "specificity_ci_upper": spe_hi, "specificity_boot_n": spe_n,
    }

# ---------------------------
# Evaluation
# ---------------------------

def evaluate_model(model, X: np.ndarray, y: np.ndarray, split: str = "train") -> Dict[str, float]:
    if X.size == 0:
        return {"accuracy": np.nan, "f1_macro": np.nan, "report": "No samples",
                "auroc": np.nan, "roc_curve": None, "roc_curves": None,
                "auroc_bootstrap": None, "sensitivity": np.nan, "specificity": np.nan,
                "accuracy_bootstrap": None, "sensitivity_bootstrap": None, "specificity_bootstrap": None}
    preds = model.predict(X)
    out = {
        "accuracy": accuracy_score(y, preds),
        "f1_macro": f1_score(y, preds, average="macro", zero_division=0),
        "report": classification_report(y, preds, zero_division=0),
        "auroc": np.nan, "roc_curve": None, "roc_curves": None, "auroc_bootstrap": None,
        "accuracy_bootstrap": None, "sensitivity_bootstrap": None, "specificity_bootstrap": None,
    }
    ovr = compute_ovr_sensitivity_specificity(y, preds, model.classes_)
    out["sensitivity"] = ovr["sensitivity"]; out["specificity"] = ovr["specificity"]

    # 테스트셋에서만 부트스트랩 CI 계산
    if split == "test" and len(np.unique(y)) > 1:
        out["accuracy_bootstrap"] = bootstrap_metric_cis(y, preds, model.classes_)

    if len(np.unique(y)) > 1:
        proba = model.predict_proba(X); classes = model.classes_
        if proba.shape[1] == 2:
            # 주의: classes[1]이 양성 레이블이라는 전제. 필요 시 명시적 매핑 도입.
            pos_scores = proba[:, 1]
            fpr, tpr, _ = roc_curve(y, pos_scores, pos_label=classes[1])
            roc_auc = auc(fpr, tpr)
            out["auroc"] = roc_auc
            out["roc_curve"] = {"fpr": fpr, "tpr": tpr, "label": str(classes[1]), "auc": roc_auc}
            y_bin = (y == classes[1]).astype(int)
            out["auroc_bootstrap"] = bootstrap_binary_auc(y_bin, pos_scores)
        else:
            curves, class_aucs = [], []
            for j, c in enumerate(classes):
                yb = (y == c).astype(int)
                if yb.min() == yb.max(): continue
                fpr, tpr, _ = roc_curve(yb, proba[:, j])
                a = auc(fpr, tpr); curves.append({"class": c, "fpr": fpr, "tpr": tpr, "auc": a}); class_aucs.append(a)
            if curves:
                macro_auc = float(np.mean(class_aucs))
                all_fpr = np.unique(np.concatenate([c["fpr"] for c in curves]))
                mean_tpr = np.zeros_like(all_fpr)
                for c in curves: mean_tpr += np.interp(all_fpr, c["fpr"], c["tpr"])
                mean_tpr /= len(curves)
                out["auroc"] = macro_auc
                out["roc_curve"] = {"fpr": all_fpr, "tpr": mean_tpr, "label": "macro", "auc": macro_auc}
                out["roc_curves"] = curves
                out["auroc_bootstrap"] = bootstrap_macro_auc(y, proba, classes)
    return out

# ---------------------------
# Persistence
# ---------------------------

def persist_model(model, feature_tag: str, feature_columns: List[str], target: str, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    model_path = Path(output_dir) / f"{feature_tag.replace('/', '-')}_{target}_random_forest.joblib"
    joblib.dump({"model": model, "feature_columns": feature_columns,
                 "target": target, "feature_tag": feature_tag,
                 "model_path": str(model_path)}, model_path)
    return model_path

def load_model(model_path: Path):
    return joblib.load(model_path)

def predict_on_full_dataset(df: pd.DataFrame, feature_cols: List[str], 
                           model_path: Path, target: str) -> pd.DataFrame:
    """전체 데이터셋에 대해 모델 예측을 수행하고 결과를 반환합니다."""
    
    # 모델 로드
    model_data = load_model(model_path)
    model = model_data["model"]
    
    # 예측할 데이터 준비 (결측값이 있는 행 제외)
    df_pred = df.dropna(subset=feature_cols).copy()
    
    if df_pred.empty:
        print(f"No valid data for prediction on {target}")
        return df
    
    # 특성 데이터 추출
    X_pred = df_pred[feature_cols].to_numpy(dtype=np.float32)
    
    # 예측 수행
    predictions = model.predict(X_pred)
    
    # 예측 결과를 원본 데이터프레임에 추가
    df_result = df.copy()
    df_result[f"{target}_pred"] = np.nan  # 기본값을 NaN으로 설정
    
    # 예측 가능한 행에만 예측 결과 할당
    df_result.loc[df_pred.index, f"{target}_pred"] = predictions
    
    return df_result

def save_predictions_with_original_data(df: pd.DataFrame, output_dir: Path) -> None:
    """예측 결과가 포함된 원본 데이터를 저장합니다."""
    
    # 예측 컬럼이 있는지 확인
    pred_cols = [col for col in df.columns if col.endswith('_pred')]
    if not pred_cols:
        print("No prediction columns found to save")
        return
    
    # CSV 파일로 저장
    output_path = output_dir / "data_with_predictions.csv"
    df.to_csv(output_path, index=False)
    print(f"Saved data with predictions to {output_path}")
    
    # 예측 결과만 별도로 저장
    pred_data = df[['phase'] + pred_cols].copy()
    pred_only_path = output_dir / "predictions_only.csv"
    pred_data.to_csv(pred_only_path, index=False)
    print(f"Saved predictions only to {pred_only_path}")

# ---------------------------
# Pipeline
# ---------------------------

def run_pipeline(args: argparse.Namespace) -> None:
    df = load_dataset(args.input_csv)
    train_df = split_by_phase(df, "train")
    eval_dfs = {"test": split_by_phase(df, "test")}
    print(f"Train set size: {len(train_df)}")
    print(f"Test set size: {len(eval_dfs['test'])}")

    feature_sets = build_feature_sets(df, methods=args.methods)
    if not feature_sets:
        raise ValueError("No quantization ratio features found for the requested methods.")

    metrics_records = []
    roc_plot_tasks = []
    shap_results = []  # SHAP 분석 결과 저장
    # 전체 데이터셋 예측을 위한 저장: feature_set -> target -> model_path
    model_paths: Dict[str, Dict[str, Path]] = {}
    # p-value 계산용 저장: target -> method -> {y_true, proba, classes}
    model_results: Dict[str, Dict[str, Dict[str, object]]] = {}

    for feature_tag, feature_cols in feature_sets.items():
        print(f"\n=== Feature set: {feature_tag} ===")
        for target in TARGETS.keys():
            print(f"-- Target: {target} --")
            t_train = prepare_target_data(train_df, target)
            if t_train.empty:
                print("No training data available. Skipping."); continue

            X_train = t_train[feature_cols].to_numpy(dtype=np.float32)
            y_train = t_train[target].to_numpy()
            model = train_model(X_train, y_train)
            model_path = persist_model(model, feature_tag, feature_cols, target, Path(args.model_dir))
            print(f"Saved model to {model_path}")
            
            # 모델 경로 저장 (전체 예측용)
            model_paths.setdefault(feature_tag, {})[target] = model_path
            
            loaded = load_model(model_path)["model"]

            # Train metrics
            m_tr = evaluate_model(loaded, X_train, y_train, split="train")
            auroc_ci_lower_tr = np.nan; auroc_ci_upper_tr = np.nan
            if m_tr.get("auroc_bootstrap"):
                auroc_ci_lower_tr = m_tr["auroc_bootstrap"].get("ci_lower", np.nan)
                auroc_ci_upper_tr = m_tr["auroc_bootstrap"].get("ci_upper", np.nan)
            metrics_records.append({"feature_set": feature_tag, "target": target, "split": "train",
                                    "accuracy": m_tr["accuracy"], "f1_macro": m_tr["f1_macro"],
                                    "auroc": m_tr.get("auroc"), "sensitivity": m_tr.get("sensitivity"),
                                    "specificity": m_tr.get("specificity"), "model_path": str(model_path),
                                    "auroc_ci_lower": auroc_ci_lower_tr, "auroc_ci_upper": auroc_ci_upper_tr})
            if m_tr.get("roc_curve") is not None:
                roc_plot_tasks.append({"feature_set": feature_tag, "target": target, "split": "train",
                                       "roc_curve": m_tr["roc_curve"], "roc_curves": m_tr.get("roc_curves")})

            # Test metrics
            for split_name, split_df in eval_dfs.items():
                t_eval = prepare_target_data(split_df, target)
                if t_eval.empty:
                    print(f"No {split_name} data available."); continue
                X_eval = t_eval[feature_cols].to_numpy(dtype=np.float32)
                y_eval = t_eval[target].to_numpy()
                m_te = evaluate_model(loaded, X_eval, y_eval, split=split_name)
                auroc_ci_lower = np.nan; auroc_ci_upper = np.nan
                if m_te.get("auroc_bootstrap"):
                    auroc_ci_lower = m_te["auroc_bootstrap"].get("ci_lower", np.nan)
                    auroc_ci_upper = m_te["auroc_bootstrap"].get("ci_upper", np.nan)
                
                # 새로운 부트스트랩 CI 필드들 추가
                accuracy_ci_lower = np.nan; accuracy_ci_upper = np.nan
                sensitivity_ci_lower = np.nan; sensitivity_ci_upper = np.nan
                specificity_ci_lower = np.nan; specificity_ci_upper = np.nan
                if m_te.get("accuracy_bootstrap"):
                    accuracy_ci_lower = m_te["accuracy_bootstrap"].get("accuracy_ci_lower", np.nan)
                    accuracy_ci_upper = m_te["accuracy_bootstrap"].get("accuracy_ci_upper", np.nan)
                    sensitivity_ci_lower = m_te["accuracy_bootstrap"].get("sensitivity_ci_lower", np.nan)
                    sensitivity_ci_upper = m_te["accuracy_bootstrap"].get("sensitivity_ci_upper", np.nan)
                    specificity_ci_lower = m_te["accuracy_bootstrap"].get("specificity_ci_lower", np.nan)
                    specificity_ci_upper = m_te["accuracy_bootstrap"].get("specificity_ci_upper", np.nan)
                
                metrics_records.append({"feature_set": feature_tag, "target": target, "split": split_name,
                                        "accuracy": m_te["accuracy"], "f1_macro": m_te["f1_macro"],
                                        "auroc": m_te.get("auroc"), "sensitivity": m_te.get("sensitivity"),
                                        "specificity": m_te.get("specificity"), "model_path": str(model_path),
                                        "auroc_ci_lower": auroc_ci_lower, "auroc_ci_upper": auroc_ci_upper,
                                        "accuracy_ci_lower": accuracy_ci_lower, "accuracy_ci_upper": accuracy_ci_upper,
                                        "sensitivity_ci_lower": sensitivity_ci_lower, "sensitivity_ci_upper": sensitivity_ci_upper,
                                        "specificity_ci_lower": specificity_ci_lower, "specificity_ci_upper": specificity_ci_upper})
                if m_te.get("roc_curve") is not None:
                    roc_plot_tasks.append({"feature_set": feature_tag, "target": target, "split": split_name,
                                           "roc_curve": m_te["roc_curve"], "roc_curves": m_te.get("roc_curves")})

                # 저장: p-value 계산 입력(테스트셋만)
                if split_name == "test":
                    proba_eval = loaded.predict_proba(X_eval)
                    classes = loaded.classes_
                    model_results.setdefault(target, {})[feature_tag] = {
                        "y_true": y_eval, "proba": proba_eval, "classes": classes
                    }
                    
    # ---------------------------
    # p-values (Binary: DeLong, Multiclass: OvR-stack DeLong with pairwise class filter)
    # ---------------------------
    print("\n=== Calculating DeLong p-values ===")
    pv_map: Dict[str, Dict[str, float]] = {}

    for target, mres in model_results.items():
        methods = mres.keys()
        if not (("dVAE" in methods) and ("AE" in methods or "metric" in methods)):
            continue

        y_true_any = mres[next(iter(methods))]["y_true"]
        is_binary = (len(np.unique(y_true_any)) == 2)

        def _p_vs(m1, m2):
            if (m1 not in mres) or (m2 not in mres):
                return np.nan
            y = mres[m1]["y_true"]
            P1 = np.asarray(mres[m1]["proba"]); P2 = np.asarray(mres[m2]["proba"])
            classes = mres[m1]["classes"]
            if is_binary:
                # classes[1]을 양성으로 사용
                return delong_pvalue_correlated_binary((y == classes[1]).astype(int), P1[:, 1], P2[:, 1])
            else:
                return delong_pvalue_correlated_multiclass(y, P1, P2, classes)

        p_dvae_ae = _p_vs("dVAE", "AE")
        p_dvae_metric = _p_vs("dVAE", "metric")
        pv_map[target] = {"AE": p_dvae_ae, "metric": p_dvae_metric}

        kind = "binary" if is_binary else "multiclass-OvR-stack"
        pa = p_dvae_ae if p_dvae_ae == p_dvae_ae else "N/A"
        pm = p_dvae_metric if p_dvae_metric == p_dvae_metric else "N/A"
        print(f"{target}: [{kind}] dVAE vs AE p={pa}; dVAE vs metric p={pm}")

    # 주입: AE/metric 테스트 행에만
    for rec in metrics_records:
        if rec["split"] != "test": continue
        tgt, meth = rec["target"], rec["feature_set"]
        if tgt in pv_map:
            if meth == "AE":
                rec["p_value_dvae_vs_ae"] = pv_map[tgt].get("AE", np.nan)
            elif meth == "metric":
                rec["p_value_dvae_vs_metric"] = pv_map[tgt].get("metric", np.nan)

    # ---------------------------
    # Outputs
    # ---------------------------
    Path(args.output_dir).mkdir(parents=True, exist_ok=True)
    
    if metrics_records:
        metrics_df = pd.DataFrame(metrics_records)
        
        # 전체 결과 저장 (기존 방식 유지)
        csv_path = Path(args.output_dir) / "classification_metrics.csv"
        metrics_df.to_csv(csv_path, index=False)
        print(f"Saved overall metrics CSV to {csv_path}")
        save_metrics_word(metrics_df, Path(args.output_dir) / "classification_metrics.docx")
        
        # 모델별 결과 저장
        print("\n=== Saving model-specific results ===")
        save_model_specific_results(metrics_df, Path(args.output_dir))

    
    # 전체 데이터셋에 대한 예측 수행
    print("\n=== Performing predictions on full dataset ===")
    df_with_predictions = df.copy()
    
    for feature_tag, target_models in model_paths.items():
        print(f"\n--- Predicting with {feature_tag} models ---")
        feature_cols = feature_sets[feature_tag]
        
        for target, model_path in target_models.items():
            print(f"Predicting {target} with {feature_tag}...")
            try:
                df_with_predictions = predict_on_full_dataset(
                    df_with_predictions, feature_cols, model_path, target
                )
                print(f"Completed prediction for {target} ({feature_tag})")
            except Exception as e:
                print(f"Prediction failed for {target} ({feature_tag}): {e}")
                continue
    
    # 예측 결과 저장
    print("\n=== Saving prediction results ===")
    save_predictions_with_original_data(df_with_predictions, Path(args.output_dir))

# ---------------------------
# CLI
# ---------------------------

def parse_arguments() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Train and evaluate models for multiple targets.")
    p.add_argument("--input-csv",
                   required=True,
                   help="Path to the dataset with quantization ratios and labels "
                        "(patient-level rows). See compute_ratios.py for schema.")
    p.add_argument("--model-dir", default="./models", help="Directory to store trained model artifacts.")
    p.add_argument("--output-dir", default="./reports", help="Directory to save reports and plots.")
    p.add_argument("--methods", nargs="+", default=METHODS,
                   help="Quantization methods to include (default: AE dVAE metric)")
    return p.parse_args()

def main() -> None:
    args = parse_arguments()
    run_pipeline(args)

if __name__ == "__main__":
    main()
